# -*- coding:utf-8 -*-
# memorandum
# MariaDB_operator.py
# author: K

import os
import re
import sys
import json
import time
from docx import Document
from win32com.client import Dispatch, constants, gencache
from sqlalchemy import create_engine, Column, Integer, String, ForeignKey  
from sqlalchemy.orm import sessionmaker, relationship  
from sqlalchemy.ext.declarative import declarative_base 
from sqlalchemy_utils import database_exists, create_database 
from . import log_ctrl, time_converter, email_ctrl, connect_database

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


class Operator:
    """The main program calls the method from here"""
    def __init__(self, id, name):
        self.data_oper = connect_database.DataOperator(id, name)
        self.memo_list = []
        self.user_name = name 
        self.id = id 

        self.logger_login = log_ctrl.memo_log(f'{name}-LOG', os.path.join(BASE_DIR, 'log', f'{name}.log'))
        self.logger_login.info(f'{name} has been login')

        self.logger_show = log_ctrl.memo_log('SHOW-LOG', os.path.join(BASE_DIR, 'log', 'show.log'))
        self.logger_add = log_ctrl.memo_log('ADD-LOG', os.path.join(BASE_DIR, 'log', 'add.log'))
        self.logger_del = log_ctrl.memo_log('DELETER-LOG', os.path.join(BASE_DIR, 'log', 'delete.log'))
        self.logger_modify = log_ctrl.memo_log('MODIFY-LOG', os.path.join(BASE_DIR, 'log', 'modify.log'))
        self.logger_query = log_ctrl.memo_log('QUERY-LOG', os.path.join(BASE_DIR, 'log', 'query.log'))
        self.logger_changetime = log_ctrl.memo_log('TIMECHANGE-LOG', os.path.join(BASE_DIR, 'log', 'timechange.log'))
        self.logger_export = log_ctrl.memo_log('EXPORT-LOG', os.path.join(BASE_DIR, 'log', 'export.log'))
        self.logger_email = log_ctrl.memo_log('EMAIL-LOG', os.path.join(BASE_DIR, 'log', 'send_email.log'))

    def add_record(self):
        """Add a memo"""
        try:
            content = input('Please enter your memo:')
            alarm_date = input('Please enter the alarm time:')
            if self.data_oper.add_record(content=content, alarm_date=alarm_date, logger=self.logger_add):
                return True
            else:
                return False
        except Exception as e:
            print(e)
            return False

    def delete_record(self):
        """Delete a memo"""
        self.data_oper.show_records(self.logger_show)
        thing_id = input('Please enter the delete_id:')
        return self.data_oper.delete_record(thing_id, self.logger_del)

    def show_record(self):
        """Display all memos"""
        return self.data_oper.show_records(self.logger_show)

    def modify_record(self):
        """Modify a memo"""
        self.show_record()
        thing_id = input('please enter memo id!')
        modify_key = input('Please enter the item of the memo：[content] or [alarm_date]:')
        modify_value = input('Please rnter the change content:')
        return self.data_oper.modify_record(thing_id, modify_key, modify_value, self.logger_modify)
    
    def export2word(self):
        """Export the information as a word file and return the file address if successful"""
        try:
            text_list = self.data_oper.return_memos()
            data_doc = Document()  
            data_doc.add_paragraph(self.user_name, 'Title')
            for m in text_list:
                p = data_doc.add_paragraph(m)
            doc_file = os.path.join(BASE_DIR, 'db', f'{self.user_name}.docx')
            data_doc.save(doc_file)
            print('word generation successful')
            self.logger_export.info(f'{self.user_name} has generation a word seccessful')
            return doc_file
        except Exception as e:
            print('WORD generation failed:', e)
            self.logger_export.error(f'{self.user_name} has generation a word failed')
            return False

    def send_email(self):
        """Send a email"""
        try:
            mail = email_ctrl.MailMaster(self.id, password=email_ctrl.get_password(), name=self.user_name)
            mail.read_email_recipients()
            file_path = self.export2word()
            mail.send_email_all('memorandum', 'There have some memos in here', attachment=file_path)
            self.logger_email.info(f'{self.user_name} has send a email successful')
            print('sending email successful')
        except Exception as e:
            self.logger_email.error(f'{self.user_name} has send a email failed')
            print('sending email failed:', e)
